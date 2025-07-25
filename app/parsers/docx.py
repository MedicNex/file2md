from .base import BaseParser
from loguru import logger
import docx
from markdownify import markdownify
import os
from PIL import Image
from app.vision import image_to_markdown, get_ocr_text, vision_client
import base64
import asyncio
from app.config import config

class DocxParser(BaseParser):
    """DOCX文件解析器"""
    
    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        return ['.docx']
    
    async def _process_image_concurrent(self, image_data: bytes, img_name: str, image_counter: int):
        """并发处理单个图片的OCR和视觉识别"""
        try:
            # 创建临时图片文件
            temp_img_path = self.create_temp_file(suffix='.png')
            
            # 保存图片
            with open(temp_img_path, 'wb') as img_file:
                img_file.write(image_data)
            
            # 并发执行OCR和视觉识别，添加异常保护
            try:
                ocr_task = get_ocr_text(temp_img_path)
                vision_task = self._get_vision_description(temp_img_path)
                
                ocr_text, vision_description = await asyncio.gather(ocr_task, vision_task)
            except Exception as ocr_vision_error:
                logger.warning(f"OCR/视觉识别失败，跳过处理 图片{image_counter}: {ocr_vision_error}")
                ocr_text = "OCR处理失败"
                vision_description = "视觉识别失败"
            
            # 生成HTML标签格式
            alt_text = f"# OCR: {ocr_text} # Visual_Features: {vision_description}"
            html_img_tag = f'<img src="{img_name}" alt="{alt_text}" />'
            
            logger.info(f"成功处理DOCX图片 {image_counter}: {img_name}")
            return f"### 图片 {image_counter}\n\n{html_img_tag}"
            
        except Exception as e:
            logger.warning(f"DOCX图片提取失败，跳过该图片 图片{image_counter}: {e}")
            return f"### 图片 {image_counter}\n\n*图片提取失败，已跳过*"
    
    async def _get_vision_description(self, temp_img_path: str) -> str:
        """获取视觉模型描述"""
        if not vision_client:
            return "视觉模型未配置"
        
        try:
            # 读取图片并转换为base64
            with open(temp_img_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            response = vision_client.chat.completions.create(
                model=os.getenv("VISION_MODEL", "gpt-4o-mini"),
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                }
                            },
                            {
                                "type": "text", 
                                "text": "Please provide a detailed description of this image, including: 1. Overall accurate description of the image; 2. Main elements and structure of the image; 3. If there are tables, charts, etc., please describe their content and layout in detail."
                            }
                        ]
                    }
                ],
                max_tokens=1000
            )
            return response.choices[0].message.content or "视觉模型识别失败"
        except Exception as e:
            logger.warning(f"Vision API调用失败: {e}")
            return "视觉模型识别失败"

    async def parse(self, file_path: str) -> str:
        """解析DOCX文件"""
        try:
            doc = docx.Document(file_path)
            
            # 提取所有段落文本
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 处理不同的段落样式
                    text = paragraph.text.strip()
                    
                    # 检查是否是标题
                    if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1])
                        text = f"{'#' * level} {text}"
                    
                    content_parts.append(text)
            
            # 处理表格
            for table in doc.tables:
                content_parts.append(self._parse_table(table))
            
            # 获取图片数量，添加保护机制
            rels = doc.part.rels
            image_count = sum(1 for rel in rels.values() if "image" in rel.target_ref)
            
            # 图片数量保护机制
            max_imgs = config.MAX_IMAGES_PER_DOC
            if max_imgs != -1 and image_count > max_imgs:
                logger.warning(f"DOCX文档包含 {image_count} 张图片，超过{max_imgs}张限制，跳过所有图片处理")
                if image_count > 0:
                    content_parts.append(f"### 文档包含 {image_count} 张图片")
                    content_parts.append(f"*因图片数量超过{max_imgs}张限制，已跳过所有图片处理*")
            else:
                # 提取并处理图片
                image_parts = await self._extract_images(doc, file_path)
                content_parts.extend(image_parts)
            
            raw_content = '\n\n'.join(content_parts)
            
            # 处理换行符
            raw_content = raw_content.replace('\\n', '\n')
            
            # 将代码块转换为HTML标签
            raw_content = self.convert_code_blocks_to_html(raw_content)
            
            # 格式化为统一的代码块格式
            markdown_content = f"```document\n{raw_content}\n```"
            
            skip_images = (config.MAX_IMAGES_PER_DOC != -1 and image_count > config.MAX_IMAGES_PER_DOC)
            logger.info(f"成功解析DOCX文件: {file_path} (总图片数: {image_count}, 跳过图片: {skip_images})")
            
            return markdown_content
            
        except Exception as e:
            logger.error(f"解析DOCX文件失败 {file_path}: {e}")
            raise Exception(f"DOCX文件解析错误: {str(e)}")
    
    def _parse_table(self, table) -> str:
        """解析表格为Markdown格式"""
        try:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    # 表头
                    rows.append('| ' + ' | '.join(cells) + ' |')
                    rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
                else:
                    # 数据行
                    rows.append('| ' + ' | '.join(cells) + ' |')
            
            return '\n'.join(rows)
        except Exception as e:
            logger.warning(f"表格解析失败: {e}")
            return "[表格解析失败]"
    
    async def _extract_images(self, doc, file_path: str) -> list[str]:
        """提取DOCX文档中的图片并进行OCR+视觉识别"""
        try:
            # 获取文档的关系部分来访问图片
            rels = doc.part.rels
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            
            # 收集所有图片信息
            image_data_list = []
            image_counter = 0
            
            for rel_id, rel in rels.items():
                if "image" in rel.target_ref:
                    image_counter += 1
                    try:
                        # 获取图片数据
                        image_data = rel.target_part.blob
                        # 生成图片文件名
                        img_name = f"{base_name}_image_{image_counter}.png"
                        
                        image_data_list.append({
                            'data': image_data,
                            'name': img_name,
                            'counter': image_counter
                        })
                    except Exception as extract_error:
                        logger.warning(f"DOCX图片提取失败，跳过该图片 图片{image_counter}: {extract_error}")
                        image_data_list.append({
                            'data': None,
                            'name': f"{base_name}_image_{image_counter}.png",
                            'counter': image_counter
                        })
            
            # 并发处理所有图片
            if image_data_list:
                image_tasks = []
                for img_info in image_data_list:
                    if img_info['data'] is not None:
                        task = self._process_image_concurrent(
                            img_info['data'], 
                            img_info['name'], 
                            img_info['counter']
                        )
                    else:
                        # 图片提取失败的情况
                        async def create_error_result(counter):
                            return f"### 图片 {counter}\n\n*图片提取失败，已跳过*"
                        task = create_error_result(img_info['counter'])
                    
                    image_tasks.append(task)
                
                try:
                    image_parts = await asyncio.gather(*image_tasks, return_exceptions=True)
                    
                    # 处理结果
                    final_image_parts = []
                    for i, result in enumerate(image_parts):
                        if isinstance(result, Exception):
                            logger.error(f"DOCX图片并发处理异常，跳过该图片 图片{i+1}: {result}")
                            final_image_parts.append(f"### 图片 {i+1}\n\n*图片处理异常，已跳过*")
                        else:
                            final_image_parts.append(result)
                    
                    if image_counter > 0:
                        logger.info(f"DOCX文档中共提取到 {image_counter} 张图片")
                    
                    return final_image_parts
                    
                except Exception as e:
                    logger.error(f"DOCX图片并发处理失败，跳过所有图片: {e}")
                    # 回退处理
                    return [f"### 图片 {i+1}\n\n*图片处理失败，已跳过*" for i in range(len(image_data_list))]
            
            return []
            
        except Exception as e:
            logger.warning(f"DOCX图片提取过程失败，跳过所有图片: {e}")
            return [] 